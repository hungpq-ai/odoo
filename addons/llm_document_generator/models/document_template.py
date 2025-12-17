import base64
import io
import logging

from odoo import api, fields, models
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)


def extract_text_from_docx(file_content):
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except ImportError:
        raise UserError("python-docx chưa được cài đặt. Chạy: pip install python-docx")
    except Exception as e:
        raise UserError(f"Lỗi đọc file DOCX: {e}")


def extract_text_from_pdf(file_content):
    """Extract text from PDF file"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_content, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n\n".join(text_parts)
    except ImportError:
        try:
            # Fallback to pdfplumber
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                return "\n\n".join(text_parts)
        except ImportError:
            raise UserError("PyMuPDF hoặc pdfplumber chưa được cài đặt. Chạy: pip install pymupdf hoặc pip install pdfplumber")
    except Exception as e:
        raise UserError(f"Lỗi đọc file PDF: {e}")


class DocumentTemplate(models.Model):
    _name = "llm.document.template"
    _description = "Document Template"
    _inherit = ["mail.thread", "mail.activity.mixin"]
    _order = "sequence, name"

    name = fields.Char(string="Tên mẫu", required=True, tracking=True, translate=True)
    description = fields.Text(string="Mô tả", translate=True)
    category_id = fields.Many2one(
        "llm.document.category",
        string="Thể loại",
        required=True,
        tracking=True,
    )
    sequence = fields.Integer(string="Thứ tự", default=10)
    active = fields.Boolean(default=True, tracking=True)
    code = fields.Char(string="Mã", help="Mã kỹ thuật cho mẫu")

    # Sample template file - user can upload
    sample_file = fields.Binary(
        string="File mẫu",
        help="Upload file mẫu (DOCX, PDF, TXT) để AI tham khảo cấu trúc",
    )
    sample_filename = fields.Char(string="Tên file")
    sample_content = fields.Text(
        string="Nội dung mẫu",
        help="Nội dung văn bản mẫu - có thể nhập trực tiếp hoặc tự động trích xuất từ file",
    )

    # Knowledge embedding for RAG
    knowledge_collection_id = fields.Many2one(
        "llm.knowledge.collection",
        string="Knowledge Collection",
        help="Collection để lưu embedded template. Giúp tiết kiệm tokens khi generate.",
    )
    knowledge_resource_id = fields.Many2one(
        "llm.resource",
        string="Template Resource",
        readonly=True,
        help="Resource đã được tạo từ template này",
    )
    is_embedded = fields.Boolean(
        string="Đã embed",
        compute="_compute_is_embedded",
        store=True,
    )

    @api.depends("knowledge_resource_id", "knowledge_resource_id.state")
    def _compute_is_embedded(self):
        for template in self:
            template.is_embedded = (
                template.knowledge_resource_id
                and template.knowledge_resource_id.state == "ready"
            )

    @api.onchange("sample_file")
    def _onchange_sample_file(self):
        """Auto-extract content when file is uploaded"""
        if not self.sample_file:
            return

        if not self.sample_filename:
            return

        try:
            # Decode file content
            file_content = base64.b64decode(self.sample_file)
            filename_lower = self.sample_filename.lower()

            if filename_lower.endswith(".docx"):
                extracted_text = extract_text_from_docx(file_content)
            elif filename_lower.endswith(".pdf"):
                extracted_text = extract_text_from_pdf(file_content)
            elif filename_lower.endswith(".txt") or filename_lower.endswith(".md"):
                extracted_text = file_content.decode("utf-8", errors="ignore")
            else:
                return {
                    "warning": {
                        "title": "Định dạng không hỗ trợ",
                        "message": f"File {self.sample_filename} không được hỗ trợ. Chỉ hỗ trợ: DOCX, PDF, TXT, MD",
                    }
                }

            if extracted_text:
                self.sample_content = extracted_text.strip()
                return {
                    "warning": {
                        "title": "Trích xuất thành công",
                        "message": f"Đã trích xuất {len(extracted_text)} ký tự từ {self.sample_filename}",
                        "type": "notification",
                    }
                }
        except UserError:
            raise
        except Exception as e:
            _logger.error(f"Error extracting file content: {e}")
            return {
                "warning": {
                    "title": "Lỗi trích xuất",
                    "message": f"Không thể trích xuất nội dung từ file: {e}",
                }
            }

    # Prompt configuration
    system_prompt = fields.Text(
        string="System Prompt",
        required=True,
        help="Hướng dẫn cho AI về vai trò và cách viết",
        default="""Bạn là chuyên gia soạn thảo văn bản chuyên nghiệp.

NHIỆM VỤ: Tạo văn bản MỚI dựa trên MẪU được cung cấp.

QUY TẮC BẮT BUỘC:
1. GIỮ NGUYÊN CẤU TRÚC của mẫu (heading, sections, thứ tự)
2. GIỮ NGUYÊN FORMAT:
   - Nếu mẫu có TABLE → tạo TABLE với cùng cấu trúc cột
   - Nếu mẫu có danh sách đánh số → giữ danh sách đánh số
   - Nếu mẫu có bullet points → giữ bullet points
3. THAY THẾ NỘI DUNG phù hợp với yêu cầu mới
4. Sử dụng Markdown format:
   - Table: | Cột 1 | Cột 2 |
   - Heading: # ## ###
   - Bold: **text**
   - List: - item hoặc 1. item
5. KHÔNG sử dụng:
   - Mermaid diagrams (```mermaid)
   - Code blocks cho sơ đồ
   - Thay vào đó, mô tả sơ đồ bằng văn bản hoặc bảng""",
    )
    user_prompt_template = fields.Text(
        string="User Prompt",
        required=True,
        help="Mẫu prompt. Sử dụng {context} cho yêu cầu user, {sample} cho nội dung mẫu",
        default="""## MẪU VĂN BẢN GỐC:

{sample}

---

## YÊU CẦU TẠO VĂN BẢN MỚI:

{context}

---

Hãy tạo văn bản mới GIỐNG ĐÚNG CẤU TRÚC mẫu trên, thay nội dung theo yêu cầu.
Nếu mẫu có bảng (table) thì văn bản mới PHẢI có bảng với cùng format.
Output bằng Markdown.""",
    )

    # Statistics
    generation_count = fields.Integer(compute="_compute_generation_count")

    def _compute_generation_count(self):
        for template in self:
            template.generation_count = self.env["llm.document.generation"].search_count(
                [("template_id", "=", template.id)]
            )

    def action_view_generations(self):
        self.ensure_one()
        return {
            "type": "ir.actions.act_window",
            "name": "Generated Documents",
            "res_model": "llm.document.generation",
            "view_mode": "list,form",
            "domain": [("template_id", "=", self.id)],
            "context": {"default_template_id": self.id},
        }

    def get_formatted_prompt(self, context, requirements=""):
        """Format the user prompt with context, sample and requirements"""
        self.ensure_one()
        sample = self.sample_content or "(Không có mẫu)"
        return self.user_prompt_template.format(
            context=context,
            sample=sample,
            requirements=requirements or "",
        )

    def action_embed_template(self):
        """Embed template sample_content into knowledge collection"""
        self.ensure_one()

        if not self.sample_content:
            raise UserError("Vui lòng nhập nội dung mẫu trước khi embed.")

        if not self.knowledge_collection_id:
            raise UserError("Vui lòng chọn Knowledge Collection trước khi embed.")

        # Get or create ir.model for llm.document.template
        model = self.env["ir.model"].search([("model", "=", "llm.document.template")], limit=1)
        if not model:
            raise UserError("Không tìm thấy model llm.document.template")

        # Check if resource already exists
        if self.knowledge_resource_id:
            # Update existing resource
            self.knowledge_resource_id.write({
                "name": f"Template: {self.name}",
                "content": self.sample_content,
                "state": "parsed",  # Reset to parsed to re-chunk
            })
            resource = self.knowledge_resource_id
        else:
            # Create new resource
            resource = self.env["llm.resource"].create({
                "name": f"Template: {self.name}",
                "model_id": model.id,
                "res_id": self.id,
                "content": self.sample_content,
                "state": "parsed",
                "collection_ids": [(4, self.knowledge_collection_id.id)],
            })
            self.knowledge_resource_id = resource

        # Process resource: chunk and embed
        resource.process_resource()

        self.message_post(
            body=f"Template đã được embed vào collection '{self.knowledge_collection_id.name}'",
            message_type="notification",
        )

        return {
            "type": "ir.actions.client",
            "tag": "display_notification",
            "params": {
                "title": "Embed thành công",
                "message": f"Template đã được embed vào '{self.knowledge_collection_id.name}'",
                "type": "success",
                "sticky": False,
            },
        }

    def get_relevant_sample(self, query, limit=5):
        """Get relevant chunks from embedded template based on query

        Args:
            query: Search query (e.g., section name or requirements)
            limit: Max number of chunks to return

        Returns:
            str: Concatenated relevant chunks or full sample_content if not embedded
        """
        self.ensure_one()

        # If not embedded, return full sample
        if not self.is_embedded or not self.knowledge_collection_id:
            return self.sample_content or ""

        try:
            # Search for relevant chunks
            chunks = self.env["llm.knowledge.chunk"].search(
                [("collection_ids", "=", self.knowledge_collection_id.id)],
                limit=limit,
                vector_search_term=query,
                collection_id=self.knowledge_collection_id.id,
            )

            if chunks:
                # Concatenate chunk contents
                return "\n\n---\n\n".join(chunk.content for chunk in chunks)
            else:
                # Fallback to full sample
                return self.sample_content or ""

        except Exception as e:
            _logger.warning(f"Error searching template chunks: {e}")
            return self.sample_content or ""
