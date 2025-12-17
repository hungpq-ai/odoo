import logging

from odoo import api, fields, models
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)


class DocumentGeneration(models.Model):
    _name = "llm.document.generation"
    _description = "Generated Document"
    _inherit = ["mail.thread", "mail.activity.mixin"]
    _order = "create_date desc"

    name = fields.Char(required=True, tracking=True)
    template_id = fields.Many2one(
        "llm.document.template",
        string="Loại tài liệu",
        required=True,
        tracking=True,
        ondelete="restrict",
    )
    category_id = fields.Many2one(
        related="template_id.category_id",
        string="Thể loại",
        store=True,
    )

    # User input - this is the main content for generation
    requirements = fields.Text(
        string="Yêu cầu",
        required=True,
        help="Yêu cầu chi tiết cho tài liệu cần tạo",
    )

    # Generated content
    generated_content = fields.Html(string="Nội dung")
    generated_markdown = fields.Text(string="Markdown")

    # Status
    state = fields.Selection(
        [
            ("draft", "Nháp"),
            ("generating", "Đang tạo"),
            ("review", "Đang review"),
            ("done", "Hoàn thành"),
            ("error", "Lỗi"),
        ],
        default="draft",
        tracking=True,
    )
    error_message = fields.Text(string="Thông báo lỗi", readonly=True)

    # Review fields
    review_feedback = fields.Text(
        string="Góp ý chỉnh sửa",
        help="Nhập góp ý để AI cải thiện tài liệu",
    )
    review_count = fields.Integer(string="Số lần review", default=0)
    review_history = fields.Html(string="Lịch sử review", readonly=True, sanitize=False)

    # LLM Selection
    selected_model_id = fields.Many2one(
        "llm.model",
        string="Model AI",
        domain="[('model_use', 'in', ['chat', 'completion'])]",
        required=True,
    )

    def _call_llm(self, model, system_prompt, user_prompt):
        """Helper to call LLM and extract response content"""
        prepend_messages = []
        if system_prompt:
            prepend_messages.append({"role": "system", "content": system_prompt})
        prepend_messages.append({"role": "user", "content": user_prompt})

        empty_messages = self.env["mail.message"].browse()
        response = model.chat(
            messages=empty_messages,
            prepend_messages=prepend_messages,
            stream=False,
        )

        if isinstance(response, dict):
            return response.get("content", "")
        elif isinstance(response, str):
            return response
        elif hasattr(response, "__iter__"):
            content = ""
            for chunk in response:
                if isinstance(chunk, dict) and chunk.get("content"):
                    content += chunk["content"]
                elif isinstance(chunk, str):
                    content += chunk
            return content
        return str(response)

    def _generate_outline(self, model, template, requirements):
        """Generate document outline (list of MAJOR sections only)"""
        # Use RAG if template is embedded, otherwise use full sample
        sample_content = template.get_relevant_sample(requirements, limit=3)

        outline_prompt = f"""Dựa trên yêu cầu sau, hãy tạo OUTLINE (dàn ý) cho tài liệu.

YÊU CẦU:
{requirements}

MẪU TÀI LIỆU (tham khảo cấu trúc):
{sample_content or "(Không có mẫu)"}

QUY TẮC QUAN TRỌNG:
- CHỈ liệt kê các PHẦN CHÍNH (major sections), KHÔNG liệt kê từng mục nhỏ
- Tối đa 5-8 sections cho toàn bộ tài liệu
- Mỗi section phải là một phần lớn, bao gồm nhiều nội dung bên trong
- VÍ DỤ ĐÚNG: "1. Giới thiệu dự án" (bao gồm tên, mã, phiên bản, mục tiêu...)
- VÍ DỤ SAI: "1. Tên dự án", "2. Mã dự án", "3. Phiên bản" (quá chi tiết)

Trả về danh sách theo format:
1. [Tên phần chính 1]: [Mô tả ngắn về nội dung bao gồm]
2. [Tên phần chính 2]: [Mô tả ngắn về nội dung bao gồm]
...

CHỈ trả về danh sách 5-8 phần chính, KHÔNG viết nội dung chi tiết."""

        return self._call_llm(model, "Bạn là chuyên gia lập dàn ý tài liệu. Chỉ tạo outline với 5-8 phần chính.", outline_prompt)

    def _generate_section(self, model, template, requirements, section_info, previous_sections=""):
        """Generate content for a single section using RAG"""
        from datetime import datetime

        # Use RAG to get relevant sample for this specific section
        relevant_sample = template.get_relevant_sample(section_info, limit=3)

        # Get current date for real-time info
        today = datetime.now().strftime("%d/%m/%Y")

        section_prompt = f"""Viết NỘI DUNG CHI TIẾT cho section sau trong tài liệu.

THÔNG TIN TỔNG QUAN:
{requirements}

SECTION CẦN VIẾT:
{section_info}

{"CÁC SECTION TRƯỚC ĐÓ (để đảm bảo liên kết):" + chr(10) + previous_sections if previous_sections else ""}

MẪU THAM KHẢO (phần liên quan):
{relevant_sample or "(Không có mẫu)"}

NGÀY HIỆN TẠI: {today}

YÊU CẦU:
- Viết chi tiết, đầy đủ nội dung cho section này
- Sử dụng Markdown format
- Giữ nguyên cấu trúc format của mẫu (table, list, heading...)
- KHÔNG lặp lại nội dung từ các section trước
- Nếu có trường ngày tháng (ngày cập nhật, ngày tạo...), sử dụng ngày hiện tại: {today}
- KHI TẠO BẢNG MARKDOWN: LUÔN có dòng separator sau header, ví dụ:
  | Cột 1 | Cột 2 | Cột 3 |
  |-------|-------|-------|
  | Dữ liệu 1 | Dữ liệu 2 | Dữ liệu 3 |"""

        return self._call_llm(model, template.system_prompt, section_prompt)

    def _convert_markdown_to_html(self, markdown_content):
        """Convert markdown to styled HTML with proper table handling"""
        import markdown
        import re

        # Pre-process: convert <br> tags to proper line breaks
        processed_content = re.sub(r'<br\s*/?>', '  \n', markdown_content)

        # Ensure blank line before tables (required for markdown parser)
        processed_content = re.sub(r'([^\n])\n(\|)', r'\1\n\n\2', processed_content)

        # Convert pipe-delimited text directly to HTML tables
        def convert_pipe_tables(content):
            lines = content.split('\n')
            result = []
            table_rows = []

            def format_cell_content(cell):
                """Convert markdown formatting within cell content"""
                # Convert `code` to <code>
                cell = re.sub(r'`([^`]+)`', r'<code style="background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; font-family: monospace;">\1</code>', cell)
                # Convert **bold** to <strong>
                cell = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', cell)
                # Convert *italic* to <em>
                cell = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', cell)
                # Convert " - " to line breaks (common list separator in cells)
                cell = re.sub(r'\s+-\s+', r'<br/>- ', cell)
                # Convert newlines to <br/>
                cell = cell.replace('\n', '<br/>')
                return cell

            def flush_table():
                if not table_rows:
                    return ''
                html = '<table style="border-collapse: collapse; width: 100%; margin: 15px 0;">'
                # First row is header
                html += '<thead style="background-color: #f8f9fa;"><tr>'
                for cell in table_rows[0]:
                    formatted_cell = format_cell_content(cell)
                    html += f'<th style="border: 1px solid #dee2e6; padding: 10px 12px; text-align: left; font-weight: bold;">{formatted_cell}</th>'
                html += '</tr></thead>'
                # Rest are body rows
                if len(table_rows) > 1:
                    html += '<tbody>'
                    for row in table_rows[1:]:
                        html += '<tr>'
                        for cell in row:
                            formatted_cell = format_cell_content(cell)
                            html += f'<td style="border: 1px solid #dee2e6; padding: 10px 12px; text-align: left; vertical-align: top;">{formatted_cell}</td>'
                        html += '</tr>'
                    html += '</tbody>'
                html += '</table>'
                return html

            i = 0
            while i < len(lines):
                line = lines[i]
                stripped = line.strip()

                # Skip empty lines while in table (don't end table on empty line)
                if not stripped and table_rows:
                    i += 1
                    continue

                # Check if line is a table row (starts and ends with |)
                if stripped.startswith('|') and stripped.endswith('|'):
                    # Skip separator rows like |---|---| or |:---|:---|
                    if re.match(r'^\|[\s\-:|]+(\|[\s\-:|]+)+\|$', stripped) and '-' in stripped:
                        i += 1
                        continue
                    # Extract cell content
                    cells = [c.strip() for c in stripped[1:-1].split('|')]
                    table_rows.append(cells)
                elif stripped.startswith('|') and not stripped.endswith('|'):
                    # Line starts with | but doesn't end with | - might be multi-line cell
                    # Accumulate until we find a line ending with |
                    full_line = stripped
                    i += 1
                    while i < len(lines) and not lines[i].strip().endswith('|'):
                        full_line += ' ' + lines[i].strip()
                        i += 1
                    if i < len(lines):
                        full_line += ' ' + lines[i].strip()
                    # Now parse the full line
                    if full_line.endswith('|'):
                        cells = [c.strip() for c in full_line[1:-1].split('|')]
                        table_rows.append(cells)
                else:
                    # End of table, convert accumulated rows to HTML
                    if table_rows:
                        result.append(flush_table())
                        table_rows = []
                    result.append(line)
                i += 1

            # Handle table at end of content
            if table_rows:
                result.append(flush_table())

            return '\n'.join(result)

        # First convert pipe tables to HTML
        processed_content = convert_pipe_tables(processed_content)

        # Protect HTML tables from markdown parser by replacing them with placeholders
        import uuid
        table_placeholders = {}
        def protect_table(match):
            placeholder = f"TABLE_PLACEHOLDER_{uuid.uuid4().hex}"
            table_placeholders[placeholder] = match.group(0)
            return placeholder

        processed_content = re.sub(r'<table[^>]*>.*?</table>', protect_table, processed_content, flags=re.DOTALL)

        # Remove standalone --- lines (horizontal rules that aren't in tables)
        processed_content = re.sub(r'^\s*---+\s*$', '', processed_content, flags=re.MULTILINE)

        # Remove markdown code block markers (```markdown, ```, etc.)
        processed_content = re.sub(r'^```\w*\s*$', '', processed_content, flags=re.MULTILINE)
        processed_content = re.sub(r'^```\s*$', '', processed_content, flags=re.MULTILINE)

        # Then convert remaining markdown (headings, lists, etc.)
        md = markdown.Markdown(extensions=['fenced_code'])
        html_content = md.convert(processed_content)

        # Restore HTML tables
        for placeholder, table_html in table_placeholders.items():
            html_content = html_content.replace(placeholder, table_html)
            # Also handle case where placeholder is wrapped in <p> tags
            html_content = html_content.replace(f'<p>{placeholder}</p>', table_html)

        # Add inline styles to headings (bold)
        html_content = re.sub(
            r'<h1>',
            '<h1 style="font-weight: bold; font-size: 1.5em; margin: 20px 0 10px 0;">',
            html_content
        )
        html_content = re.sub(
            r'<h2>',
            '<h2 style="font-weight: bold; font-size: 1.3em; margin: 18px 0 8px 0;">',
            html_content
        )
        html_content = re.sub(
            r'<h3>',
            '<h3 style="font-weight: bold; font-size: 1.1em; margin: 15px 0 6px 0;">',
            html_content
        )
        html_content = re.sub(
            r'<h4>',
            '<h4 style="font-weight: bold; font-size: 1em; margin: 12px 0 5px 0;">',
            html_content
        )
        html_content = re.sub(
            r'<h5>',
            '<h5 style="font-weight: bold; font-size: 0.95em; margin: 10px 0 4px 0;">',
            html_content
        )
        html_content = re.sub(
            r'<h6>',
            '<h6 style="font-weight: bold; font-size: 0.9em; margin: 8px 0 4px 0;">',
            html_content
        )

        # Convert numbered section headers (e.g., "2. Title" or "3.2.1. Title") to bold
        # Match: start of paragraph, number pattern, text until end of line
        html_content = re.sub(
            r'<p>(\d+\.(?:\d+\.)*\s+[^<]+)</p>',
            r'<p style="font-weight: bold; font-size: 1.1em; margin: 15px 0 8px 0;">\1</p>',
            html_content
        )

        # Convert **text** to bold (markdown bold syntax that wasn't converted)
        html_content = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', html_content)

        return html_content

    def action_generate(self):
        """Generate document using LLM - section by section for long documents"""
        import re

        for record in self:
            try:
                record.state = "generating"
                record._cr.commit()

                template = record.template_id
                model = record.selected_model_id
                if not model:
                    raise UserError("Vui lòng chọn Model AI.")

                # Step 1: Generate outline
                _logger.info(f"Generating outline for document {record.id}")
                record.message_post(body="Đang tạo dàn ý...", message_type="notification")
                record._cr.commit()

                outline = record._generate_outline(model, template, record.requirements)
                if not outline:
                    raise UserError("Không thể tạo dàn ý")

                # Parse sections from outline
                sections = []
                for line in outline.split('\n'):
                    line = line.strip()
                    if re.match(r'^\d+\.?\s*', line):
                        sections.append(re.sub(r'^\d+\.?\s*', '', line))

                if not sections:
                    # Fallback: treat entire outline as one section
                    sections = [outline]

                _logger.info(f"Found {len(sections)} sections to generate")

                # Step 2: Generate each section
                all_content = []
                for i, section in enumerate(sections):
                    _logger.info(f"Generating section {i+1}/{len(sections)}: {section[:50]}...")
                    record.message_post(
                        body=f"Đang tạo phần {i+1}/{len(sections)}: {section[:50]}...",
                        message_type="notification"
                    )
                    record._cr.commit()

                    # Get previous sections summary for context
                    previous_summary = "\n".join(all_content[-2:]) if all_content else ""

                    section_content = record._generate_section(
                        model, template, record.requirements,
                        section, previous_summary
                    )

                    if section_content:
                        all_content.append(section_content)

                    # Commit after each section to save progress
                    record._cr.commit()

                # Step 3: Combine all sections
                if all_content:
                    full_markdown = "\n\n---\n\n".join(all_content)
                    record.generated_markdown = full_markdown
                    record.generated_content = record._convert_markdown_to_html(full_markdown)
                    record.state = "done"
                    record.message_post(
                        body=f"Tạo tài liệu thành công! ({len(sections)} phần)",
                        message_type="notification",
                    )
                else:
                    raise UserError("Không tạo được nội dung")

            except Exception as e:
                _logger.error(f"Error generating document {record.id}: {str(e)}")
                record.state = "error"
                record.error_message = str(e)
                record.message_post(
                    body=f"Lỗi tạo tài liệu: {str(e)}",
                    message_type="notification",
                )

    def action_regenerate(self):
        """Regenerate the document"""
        self.write({
            "state": "draft",
            "generated_content": False,
            "generated_markdown": False,
            "error_message": False,
        })
        return self.action_generate()

    def action_start_review(self):
        """Start review mode"""
        self.ensure_one()
        self.state = "review"
        self.message_post(
            body="Bắt đầu review tài liệu. Nhập góp ý và bấm 'Áp dụng góp ý' để AI cải thiện.",
            message_type="notification",
        )

    def action_apply_review(self):
        """Apply review feedback using AI"""
        from datetime import datetime

        self.ensure_one()

        if not self.review_feedback:
            raise UserError("Vui lòng nhập góp ý chỉnh sửa trước khi áp dụng.")

        if not self.generated_markdown:
            raise UserError("Không có nội dung để chỉnh sửa.")

        model = self.selected_model_id
        if not model:
            raise UserError("Vui lòng chọn Model AI.")

        try:
            self.state = "generating"
            self._cr.commit()

            today = datetime.now().strftime("%d/%m/%Y")

            review_prompt = f"""Bạn được giao nhiệm vụ CHỈNH SỬA tài liệu dựa trên góp ý của người dùng.

TÀI LIỆU HIỆN TẠI:
{self.generated_markdown}

GÓP Ý CHỈNH SỬA:
{self.review_feedback}

NGÀY HIỆN TẠI: {today}

YÊU CẦU:
1. Đọc kỹ góp ý và áp dụng các thay đổi được yêu cầu
2. GIỮ NGUYÊN các phần không được đề cập trong góp ý
3. Giữ format Markdown
4. Giữ cấu trúc bảng, heading, list như cũ
5. Nếu góp ý yêu cầu thêm/sửa nội dung, hãy thực hiện phù hợp với ngữ cảnh
6. Cập nhật ngày tháng nếu cần: {today}

Trả về TÀI LIỆU ĐÃ CHỈNH SỬA HOÀN CHỈNH (không giải thích, chỉ trả về nội dung)."""

            system_prompt = "Bạn là chuyên gia chỉnh sửa văn bản. Áp dụng chính xác các góp ý mà không làm mất nội dung khác."

            revised_content = self._call_llm(model, system_prompt, review_prompt)

            if revised_content:
                # Save review history with HTML formatting
                review_num = self.review_count + 1
                history_entry = f'''
                <div style="border: 1px solid #dee2e6; border-radius: 8px; padding: 12px; margin-bottom: 10px; background-color: #f8f9fa;">
                    <div style="font-weight: bold; color: #495057; margin-bottom: 8px;">
                        <i class="fa fa-edit" style="margin-right: 5px;"></i>Review #{review_num} - {today}
                    </div>
                    <div style="color: #212529; white-space: pre-wrap;">{self.review_feedback}</div>
                </div>
                '''
                self.review_history = (self.review_history or "") + history_entry

                # Update content
                self.generated_markdown = revised_content
                self.generated_content = self._convert_markdown_to_html(revised_content)
                self.review_count += 1
                self.review_feedback = False  # Clear feedback
                self.state = "done"

                self.message_post(
                    body=f"Đã áp dụng góp ý (lần {self.review_count}). Tài liệu đã được cập nhật.",
                    message_type="notification",
                )
            else:
                raise UserError("Không thể áp dụng góp ý. Vui lòng thử lại.")

        except Exception as e:
            _logger.error(f"Error applying review: {str(e)}")
            self.state = "review"
            raise UserError(f"Lỗi khi áp dụng góp ý: {str(e)}")

    def action_approve_document(self):
        """Approve document and finish review"""
        self.ensure_one()
        self.state = "done"
        self.review_feedback = False
        self.message_post(
            body="Tài liệu đã được duyệt.",
            message_type="notification",
        )

    def action_reset_draft(self):
        """Reset to draft state"""
        self.write({
            "state": "draft",
            "error_message": False,
        })

    def _parse_markdown_table(self, lines, start_idx):
        """Parse markdown table starting at given index, return (table_data, end_idx)"""
        table_rows = []
        idx = start_idx

        while idx < len(lines):
            line = lines[idx].strip()
            # Check if line is a table row (starts and ends with |)
            if line.startswith("|") and "|" in line[1:]:
                # Parse cells
                cells = [cell.strip() for cell in line.split("|")[1:-1]]
                # Skip separator row (contains only dashes and colons)
                if not all(c.replace("-", "").replace(":", "").strip() == "" for c in cells):
                    table_rows.append(cells)
                idx += 1
            else:
                break

        return table_rows, idx

    def _add_table_to_docx(self, doc, table_data):
        """Add a table to DOCX document with proper formatting"""
        if not table_data:
            return

        import re
        from docx.shared import Pt, Inches

        # Create table
        num_cols = max(len(row) for row in table_data)
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        table.style = "Table Grid"

        # Fill table
        for row_idx, row_data in enumerate(table_data):
            row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = row.cells[col_idx]
                    # Clear default paragraph
                    cell.text = ""
                    paragraph = cell.paragraphs[0]

                    # Process cell text with formatting
                    self._add_formatted_text(paragraph, cell_text)

                    # Bold header row
                    if row_idx == 0:
                        for run in paragraph.runs:
                            run.bold = True

    def _clean_text(self, text):
        """Clean markdown/html artifacts from text"""
        import re
        # Replace <br> and <br/> with newline
        text = re.sub(r'<br\s*/?>', '\n', text)
        # Remove ** markers (will handle bold separately)
        return text

    def _add_formatted_text(self, paragraph, text):
        """Add text with bold formatting and line breaks to paragraph"""
        import re
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # First handle <br> tags - replace with newline
        text = re.sub(r'<br\s*/?>', '\n', text)

        # Split by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                inner_text = part[2:-2]
                # Handle newlines within bold text
                lines = inner_text.split('\n')
                for i, line in enumerate(lines):
                    run = paragraph.add_run(line)
                    run.bold = True
                    if i < len(lines) - 1:
                        # Add line break
                        run.add_break()
            else:
                # Handle newlines in regular text
                lines = part.split('\n')
                for i, line in enumerate(lines):
                    paragraph.add_run(line)
                    if i < len(lines) - 1:
                        # Add line break
                        paragraph.add_run().add_break()

    def action_export_docx(self):
        """Export as DOCX file with table support"""
        self.ensure_one()
        if not self.generated_markdown:
            raise UserError("Không có nội dung để xuất")

        import base64
        import io
        import re

        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise UserError("Cần cài đặt thư viện python-docx để xuất file DOCX")

        # Create DOCX document
        doc = Document()

        # Convert markdown to docx paragraphs (no auto title)
        lines = self.generated_markdown.split("\n")
        idx = 0

        while idx < len(lines):
            line = lines[idx]
            stripped = line.strip()

            if not stripped:
                doc.add_paragraph()
                idx += 1
            # Check for table start
            elif stripped.startswith("|") and "|" in stripped[1:]:
                table_data, idx = self._parse_markdown_table(lines, idx)
                if table_data:
                    self._add_table_to_docx(doc, table_data)
                    doc.add_paragraph()  # Add space after table
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
                idx += 1
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
                idx += 1
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
                idx += 1
            elif stripped.startswith("#### "):
                doc.add_heading(stripped[5:], level=4)
                idx += 1
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                self._add_formatted_text(p, stripped[2:])
                idx += 1
            elif re.match(r'^\d+\.\s', stripped):
                # Handle numbered list (any number)
                text = re.sub(r'^\d+\.\s', '', stripped)
                p = doc.add_paragraph(style="List Number")
                self._add_formatted_text(p, text)
                idx += 1
            elif stripped.startswith("**") and stripped.endswith("**"):
                p = doc.add_paragraph()
                p.add_run(stripped[2:-2]).bold = True
                idx += 1
            else:
                p = doc.add_paragraph()
                self._add_formatted_text(p, stripped)
                idx += 1

        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Create attachment
        attachment = self.env["ir.attachment"].create({
            "name": f"{self.name}.docx",
            "type": "binary",
            "datas": base64.b64encode(file_stream.read()),
            "res_model": self._name,
            "res_id": self.id,
            "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        })

        return {
            "type": "ir.actions.act_url",
            "url": f"/web/content/{attachment.id}?download=true",
            "target": "self",
        }

    def action_export_markdown(self):
        """Export as markdown file"""
        self.ensure_one()
        if not self.generated_markdown:
            raise UserError("Không có nội dung để xuất")

        import base64

        # Create attachment
        attachment = self.env["ir.attachment"].create({
            "name": f"{self.name}.md",
            "type": "binary",
            "datas": base64.b64encode(self.generated_markdown.encode("utf-8")),
            "res_model": self._name,
            "res_id": self.id,
            "mimetype": "text/markdown",
        })

        return {
            "type": "ir.actions.act_url",
            "url": f"/web/content/{attachment.id}?download=true",
            "target": "self",
        }

    def action_export_html(self):
        """Export as HTML file"""
        self.ensure_one()
        if not self.generated_content:
            raise UserError("No content to export")

        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{self.name}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
{self.generated_content}
</body>
</html>"""

        attachment = self.env["ir.attachment"].create({
            "name": f"{self.name}.html",
            "type": "binary",
            "datas": html_content.encode("utf-8"),
            "res_model": self._name,
            "res_id": self.id,
            "mimetype": "text/html",
        })

        return {
            "type": "ir.actions.act_url",
            "url": f"/web/content/{attachment.id}?download=true",
            "target": "self",
        }

    def _get_html_for_pdf(self):
        """Get HTML content formatted for PDF export"""
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{self.name}</title>
    <style>
        @page {{
            size: A4;
            margin: 2cm;
        }}
        body {{
            font-family: 'DejaVu Sans', Arial, sans-serif;
            font-size: 12pt;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{
            font-size: 18pt;
            text-align: center;
            margin-bottom: 20px;
            color: #000;
        }}
        h2 {{
            font-size: 14pt;
            margin-top: 15px;
            margin-bottom: 10px;
            color: #000;
        }}
        h3 {{
            font-size: 12pt;
            margin-top: 12px;
            margin-bottom: 8px;
            color: #000;
        }}
        p {{
            margin: 8px 0;
            text-align: justify;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 15px 0;
        }}
        th, td {{
            border: 1px solid #333;
            padding: 8px 10px;
            text-align: left;
        }}
        th {{
            background-color: #f0f0f0;
            font-weight: bold;
        }}
        ul, ol {{
            margin: 10px 0;
            padding-left: 25px;
        }}
        li {{
            margin: 5px 0;
        }}
    </style>
</head>
<body>
    {self.generated_content}
</body>
</html>"""

    def action_export_pdf(self):
        """Export as PDF file using wkhtmltopdf"""
        self.ensure_one()
        if not self.generated_content:
            raise UserError("Không có nội dung để xuất")

        import base64
        import subprocess
        import tempfile
        import os

        html_content = self._get_html_for_pdf()

        # Create temp files
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as html_file:
            html_file.write(html_content)
            html_path = html_file.name

        pdf_path = html_path.replace('.html', '.pdf')

        try:
            # Use wkhtmltopdf (available in Odoo container)
            cmd = [
                'wkhtmltopdf',
                '--encoding', 'utf-8',
                '--page-size', 'A4',
                '--margin-top', '20mm',
                '--margin-bottom', '20mm',
                '--margin-left', '20mm',
                '--margin-right', '20mm',
                '--enable-local-file-access',
                html_path,
                pdf_path
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

            if not os.path.exists(pdf_path):
                _logger.error(f"wkhtmltopdf error: {result.stderr}")
                raise UserError(f"Lỗi tạo PDF: {result.stderr}")

            # Read PDF content
            with open(pdf_path, 'rb') as pdf_file:
                pdf_data = pdf_file.read()

            # Create attachment
            attachment = self.env["ir.attachment"].create({
                "name": f"{self.name}.pdf",
                "type": "binary",
                "datas": base64.b64encode(pdf_data),
                "res_model": self._name,
                "res_id": self.id,
                "mimetype": "application/pdf",
            })

            return {
                "type": "ir.actions.act_url",
                "url": f"/web/content/{attachment.id}?download=true",
                "target": "self",
            }

        finally:
            # Cleanup temp files
            if os.path.exists(html_path):
                os.unlink(html_path)
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
